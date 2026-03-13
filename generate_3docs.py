#!/usr/bin/env python3
"""3つのマーケティング資料を一括生成"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ===== 共通ユーティリティ =====
NAVY = RGBColor(0x0D, 0x2B, 0x4E)
DARK_BLUE = RGBColor(0x1A, 0x47, 0x7A)
MID_BLUE = RGBColor(0x2E, 0x75, 0xB6)
ACCENT_GOLD = RGBColor(0xB8, 0x86, 0x0B)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED_ACCENT = RGBColor(0xC0, 0x39, 0x2B)

OFFICE_NAME = '弁護士法人 東京新橋法律事務所'
REPRESENTATIVE = '代表弁護士 前田 祥夢'
ADDRESS = '〒104-0061 東京都中央区銀座8丁目20番33号 ACN銀座8ビル3F'
TEL = '03-6273-3254'
URL = 'https://tslaw.or.jp'
HOURS = '平日 10:00〜19:00'


def set_font(run, name='Yu Gothic', size=Pt(10.5), color=DARK_GRAY, bold=False):
    run.font.name = name
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)


def add_styled_para(doc, text, size=Pt(10.5), color=DARK_GRAY, bold=False, alignment=None, space_before=0, space_after=Pt(6), line_spacing=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold)
    p.paragraph_format.space_before = Pt(space_before) if isinstance(space_before, (int, float)) else space_before
    p.paragraph_format.space_after = space_after
    if alignment:
        p.alignment = alignment
    if line_spacing:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = line_spacing
    return p


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_font(run, size=Pt(18 if level == 1 else 14 if level == 2 else 12),
                color=NAVY if level == 1 else DARK_BLUE if level == 2 else MID_BLUE,
                bold=True)
    return h


def add_divider(doc, color='2E75B6'):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(12)
    return p


def add_styled_table(doc, headers, rows, col_widths=None, header_bg='0D2B4E'):
    from docx.shared import Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, size=Pt(9.5), color=WHITE, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), header_bg)
        shd.set(qn('w:val'), 'clear')
        shading.append(shd)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_font(run, size=Pt(9))
            if r_idx % 2 == 1:
                shading = cell._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F2F7FB')
                shd.set(qn('w:val'), 'clear')
                shading.append(shd)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_footer_info(doc):
    add_divider(doc, 'B8860B')
    p = doc.add_paragraph()
    lines = [
        (OFFICE_NAME, Pt(10), True, NAVY),
        (REPRESENTATIVE, Pt(9), False, DARK_GRAY),
        (ADDRESS, Pt(8.5), False, MED_GRAY),
        (f'TEL {TEL}（{HOURS}）', Pt(8.5), False, MED_GRAY),
        (URL, Pt(8.5), False, MID_BLUE),
    ]
    for i, (text, size, bold, color) in enumerate(lines):
        run = p.add_run(text)
        set_font(run, size=size, color=color, bold=bold)
        if i < len(lines) - 1:
            run = p.add_run('\n')
            run.font.size = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)


# =====================================================
# ① 管理会社向けDM
# =====================================================
def create_dm_for_management():
    doc = Document()

    # ページ設定（A4）
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # スタイル
    style = doc.styles['Normal']
    style.font.name = 'Yu Gothic'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

    # ===== ヘッダー部分 =====
    add_styled_para(doc, '2026年3月吉日', size=Pt(10), color=MED_GRAY, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    p = doc.add_paragraph()
    run = p.add_run('賃貸管理会社　ご担当者様')
    set_font(run, size=Pt(11), bold=True)
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    for text, size, bold, color in [
        (OFFICE_NAME, Pt(10), False, DARK_GRAY),
        ('\n', Pt(4), False, DARK_GRAY),
        (REPRESENTATIVE, Pt(10), True, DARK_GRAY),
    ]:
        run = p.add_run(text)
        set_font(run, size=size, color=color, bold=bold)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(18)

    # ===== タイトル =====
    add_divider(doc, '0D2B4E')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('高齢入居者の「孤独死リスク」を\nゼロにする仕組みのご提案')
    set_font(run, size=Pt(16), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— 国交省推奨「死後事務委任契約」の活用による、\n賃貸管理の新しいリスクヘッジ —')
    set_font(run, size=Pt(10), color=MID_BLUE)
    p.paragraph_format.space_after = Pt(6)

    add_divider(doc, '0D2B4E')

    # ===== 本文 =====
    add_styled_para(doc, '拝啓　時下ますますご清栄のこととお慶び申し上げます。突然のご連絡失礼いたします。', space_after=Pt(8))

    add_styled_para(doc, '私は東京・銀座で法律事務所を経営しております弁護士の前田祥夢と申します。当事務所は、家賃保証会社様と連携した建物明渡業務を多数取り扱っており、賃貸管理の現場で発生する法的課題に日常的に向き合っております。', space_after=Pt(12))

    # --- 問題提起 ---
    p = doc.add_paragraph()
    run = p.add_run('■ こんなお悩みはございませんか？')
    set_font(run, size=Pt(12), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(8)

    problems = [
        '高齢の単身入居者が亡くなった後、賃貸借契約の解約ができない',
        '残置物の処分に法的根拠がなく、何ヶ月も放置せざるを得ない',
        '孤独死が発生し、原状回復に数十万円〜の費用が発生した',
        '事故物件化により、家賃の大幅値下げを余儀なくされた',
        '高齢者の入居申込を断りたいが、空室率も気になる',
    ]
    for prob in problems:
        p = doc.add_paragraph()
        run = p.add_run(f'  {prob}')
        set_font(run, size=Pt(10))
        p.paragraph_format.space_after = Pt(2)
        # チェックマーク
        p.runs[0].text = f'  \u2611  {prob}'

    doc.add_paragraph()

    # --- 解決策 ---
    p = doc.add_paragraph()
    run = p.add_run('■ 解決策：「死後事務委任契約」とは')
    set_font(run, size=Pt(12), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(8)

    add_styled_para(doc, '「死後事務委任契約」とは、ご本人が元気なうちに弁護士と契約を結び、万が一の際に以下を弁護士が代行する仕組みです。')

    services = [
        ('賃貸借契約の解約', '法的権限に基づき、速やかに契約を終了'),
        ('残置物の撤去・処分', '弁護士の権限で適法に対応。数ヶ月待つ必要なし'),
        ('原状回復の手配', '遺産または預託金から費用を支出'),
        ('死亡届の提出', '行政手続きを弁護士が全て代行'),
        ('葬儀・納骨の手配', '事前の意向に沿って手配'),
        ('関係者への通知', '大家様・管理会社様への迅速な死亡連絡'),
    ]
    add_styled_table(doc,
        ['対応事項', '内容'],
        [[s[0], s[1]] for s in services],
        col_widths=[2.0, 4.5]
    )

    # --- 制度的裏付け ---
    p = doc.add_paragraph()
    run = p.add_run('■ 国の制度的裏付け')
    set_font(run, size=Pt(12), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(8)

    add_styled_para(doc, '2021年、国土交通省・法務省は「残置物の処理等に関するモデル契約条項」を策定し、単身高齢者の賃貸借において「死後事務委任契約」の活用を公式に推奨しています。', space_after=Pt(4))
    add_styled_para(doc, 'また、2024年の改正住宅セーフティネット法により、居住支援法人の業務に「残置物処理」が追加され、死後事務委任契約の社会的な必要性はますます高まっています。', space_after=Pt(12))

    # --- 管理会社のメリット ---
    p = doc.add_paragraph()
    run = p.add_run('■ 貴社のメリット')
    set_font(run, size=Pt(12), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(8)

    add_styled_table(doc,
        ['メリット', '詳細'],
        [
            ['孤独死リスクの実質ゼロ化', '弁護士が全ての死後手続きを代行。管理会社様のご負担は一切なし'],
            ['残置物問題の即時解決', '法的権限に基づく撤去。数ヶ月の放置期間が不要に'],
            ['高齢者の入居受入が可能に', '死後事務委任契約済みの入居者は「安全な入居者」に'],
            ['空室率の改善', '高齢者を受け入れることで入居対象が拡大'],
            ['費用負担ゼロ', '弁護士費用は入居者ご本人が負担。管理会社様の費用は一切不要'],
        ],
        col_widths=[2.0, 4.5]
    )

    # --- 導入の流れ ---
    p = doc.add_paragraph()
    run = p.add_run('■ 導入の流れ')
    set_font(run, size=Pt(12), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(8)

    steps = [
        ('STEP 1', '管理会社様と当事務所で提携契約を締結（費用不要）'),
        ('STEP 2', '高齢単身入居者に対し、当事務所の死後事務委任サービスをご案内'),
        ('STEP 3', '入居者ご本人と当事務所が直接契約を締結'),
        ('STEP 4', '万が一の際、当事務所が全ての死後事務を代行'),
    ]
    for step, desc in steps:
        p = doc.add_paragraph()
        run1 = p.add_run(f'  {step}  ')
        set_font(run1, size=Pt(10), color=WHITE, bold=True)
        # 背景色をrun levelでは難しいのでテキストで表現
        run1.text = f'  【{step}】 '
        set_font(run1, size=Pt(10), color=MID_BLUE, bold=True)
        run2 = p.add_run(desc)
        set_font(run2, size=Pt(10))
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    # --- お問い合わせ ---
    p = doc.add_paragraph()
    run = p.add_run('■ まずはお気軽にご連絡ください')
    set_font(run, size=Pt(12), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(8)

    add_styled_para(doc, '管理物件における高齢入居者の現状や課題について、無料でご相談をお受けしております。貴社の状況に合わせた最適な導入プランをご提案いたします。')
    add_styled_para(doc, 'まずは30分のオンライン面談で、具体的な活用方法をご説明させていただければ幸いです。')
    add_styled_para(doc, 'ご多忙のところ最後までお読みいただき、誠にありがとうございました。', space_after=Pt(2))
    add_styled_para(doc, '敬具', alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=Pt(18))

    # フッター
    add_footer_info(doc)

    doc.save(r'C:\Users\user\Desktop\claudeマスター\【管理会社向け】孤独死リスク対策DM.docx')
    print('✓ 管理会社向けDM 生成完了')


# =====================================================
# ② 新聞折込広告
# =====================================================
def create_newspaper_ad():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    style = doc.styles['Normal']
    style.font.name = 'Yu Gothic'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

    # ===== メインキャッチ =====
    doc.add_paragraph()  # 上部余白

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ひとりで暮らす あなたへ')
    set_font(run, size=Pt(28), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(4)

    add_divider(doc, 'B8860B')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('「もしも」のとき、\n誰があなたの代わりを務めますか？')
    set_font(run, size=Pt(18), color=DARK_BLUE, bold=True)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5

    # ===== 本文（大きめフォント） =====
    body_texts = [
        'おひとりで暮らしていると、ふと不安になることはありませんか。',
        '',
        '「自分に何かあったとき、届出や手続きは誰がやるのだろう」',
        '「葬儀の手配は？ 部屋の片付けは？ 誰に頼めばいいのか」',
        '',
        'こうした「もしも」の備えを、',
        '弁護士があなたに代わってお引き受けいたします。',
    ]
    for text in body_texts:
        if text == '':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
        elif text.startswith('「'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_font(run, size=Pt(13), color=DARK_BLUE)
            p.paragraph_format.space_after = Pt(4)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_font(run, size=Pt(13), color=DARK_GRAY)
            p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    # ===== サービス内容 =====
    add_divider(doc, '2E75B6')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('弁護士がお手伝いできること')
    set_font(run, size=Pt(16), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(12)

    items = [
        '\u2713  届出・行政手続き（死亡届、年金停止、保険証返却など）',
        '\u2713  葬儀・火葬・納骨の手配',
        '\u2713  お部屋の片付け・明け渡し',
        '\u2713  公共料金の解約手続き',
        '\u2713  ご関係者への連絡',
        '\u2713  その他、おひとりでは備えにくい「もしも」の手続き全般',
    ]
    for item in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        set_font(run, size=Pt(12), color=DARK_GRAY)
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()
    add_divider(doc, '2E75B6')

    # ===== CTA =====
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('まずはお電話ください（ご相談無料）')
    set_font(run, size=Pt(14), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TEL)
    set_font(run, size=Pt(28), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'受付時間：{HOURS}')
    set_font(run, size=Pt(11), color=MED_GRAY)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('「新聞を見た」とお伝えいただければスムーズです')
    set_font(run, size=Pt(10), color=MID_BLUE)
    p.paragraph_format.space_after = Pt(18)

    # ===== 事務所情報 =====
    add_divider(doc, 'B8860B')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, size, bold, color in [
        (OFFICE_NAME, Pt(12), True, NAVY),
        ('\n', Pt(2), False, DARK_GRAY),
        (REPRESENTATIVE, Pt(10), False, DARK_GRAY),
        ('\n', Pt(2), False, DARK_GRAY),
        (ADDRESS, Pt(9), False, MED_GRAY),
        ('\n', Pt(2), False, DARK_GRAY),
        (URL, Pt(9), False, MID_BLUE),
        ('\n', Pt(6), False, DARK_GRAY),
        ('東京弁護士会所属  初回相談無料', Pt(9), False, MED_GRAY),
    ]:
        run = p.add_run(text)
        set_font(run, size=size, color=color, bold=bold)

    doc.save(r'C:\Users\user\Desktop\claudeマスター\【新聞折込】おひとりさまの安心相談.docx')
    print('✓ 新聞折込広告 生成完了')


# =====================================================
# ③ 汎用リーフレット
# =====================================================
def create_leaflet():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = 'Yu Gothic'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

    # ===== 表紙（1ページ目） =====
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('おひとりさまの')
    set_font(run, size=Pt(16), color=MID_BLUE)
    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('「もしも」に備える')
    set_font(run, size=Pt(24), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('安心サポート')
    set_font(run, size=Pt(24), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(16)

    add_divider(doc, 'B8860B')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('弁護士が、あなたの「最後の届け人」になります')
    set_font(run, size=Pt(13), color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, size, bold, color in [
        (OFFICE_NAME, Pt(12), True, NAVY),
        ('\n', Pt(4), False, DARK_GRAY),
        (REPRESENTATIVE, Pt(10), False, DARK_GRAY),
    ]:
        run = p.add_run(text)
        set_font(run, size=size, color=color, bold=bold)

    doc.add_page_break()

    # ===== 2ページ目：問題提起 =====
    p = doc.add_paragraph()
    run = p.add_run('こんなお悩みはありませんか？')
    set_font(run, size=Pt(16), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(12)

    worries = [
        '自分が亡くなった後、届出や手続きをしてくれる人がいない',
        '葬儀やお墓のことを頼める家族がいない',
        '賃貸の部屋をきちんと片付けて返したいが、誰に頼めばいいかわからない',
        '銀行口座や保険の手続きが、自分の死後に宙に浮いてしまいそう',
        '「迷惑をかけたくない」が、具体的に何をすればいいのかわからない',
    ]
    for w in worries:
        p = doc.add_paragraph()
        run = p.add_run(f'  \u2611  {w}')
        set_font(run, size=Pt(11))
        p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ひとつでも当てはまる方は、ぜひお読みください。')
    set_font(run, size=Pt(12), color=DARK_BLUE, bold=True)
    p.paragraph_format.space_after = Pt(18)

    add_divider(doc, '2E75B6')

    # --- サービス説明 ---
    p = doc.add_paragraph()
    run = p.add_run('「死後事務委任」という安心の仕組み')
    set_font(run, size=Pt(14), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(8)

    add_styled_para(doc, 'お元気なうちに弁護士と契約を結んでおくことで、万が一の際、弁護士があなたに代わって必要な手続きを全て行います。', size=Pt(11), space_after=Pt(12))

    add_styled_table(doc,
        ['弁護士がお引き受けすること', '具体的な内容'],
        [
            ['届出・行政手続き', '死亡届の提出、年金停止、健康保険証の返却など'],
            ['葬儀・火葬・納骨', 'ご本人の事前のご希望に沿って手配いたします'],
            ['お部屋の明け渡し', '賃貸契約の解約、原状回復、遺品の整理'],
            ['各種解約手続き', '公共料金、携帯電話、各種サービスの解約'],
            ['関係者へのご連絡', 'ご友人・ご知人など、事前にご指定いただいた方へ'],
            ['デジタルの整理', 'SNSアカウント、メールなどの削除・整理'],
        ],
        col_widths=[2.0, 4.5]
    )

    doc.add_page_break()

    # ===== 3ページ目：料金・Q&A =====
    p = doc.add_paragraph()
    run = p.add_run('ご利用の流れ')
    set_font(run, size=Pt(14), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(12)

    flow = [
        ('1. 無料相談', 'お電話またはご来所で、お悩みやご希望をお聞かせください。\nこの段階では費用は一切かかりません。'),
        ('2. プランのご提案', 'お話を伺った上で、必要な手続きの範囲と費用をご説明します。\nご納得いただけるまで、何度でもご相談いただけます。'),
        ('3. 契約の締結', '公正証書で契約を結びます。\nこれにより、法的に確実な備えが完成します。'),
        ('4. 安心の日常', '契約後は、いつも通りの生活をお過ごしください。\n定期的なご連絡で、お元気であることを確認いたします。'),
        ('5. 万が一の際', '弁護士が全ての手続きを責任を持って行います。\n事前のご希望に沿って、ひとつひとつ丁寧に進めます。'),
    ]
    for title, desc in flow:
        p = doc.add_paragraph()
        run = p.add_run(title)
        set_font(run, size=Pt(11), color=DARK_BLUE, bold=True)
        p.paragraph_format.space_after = Pt(2)

        p = doc.add_paragraph()
        run = p.add_run(desc)
        set_font(run, size=Pt(10), color=DARK_GRAY)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Inches(0.3)

    add_divider(doc, '2E75B6')

    # 料金
    p = doc.add_paragraph()
    run = p.add_run('料金の目安')
    set_font(run, size=Pt(14), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(8)

    add_styled_table(doc,
        ['プラン', '内容', '契約時の費用', '備考'],
        [
            ['シンプル', '届出・葬儀・納骨・行政手続き', '22万円〜', '最低限の備えをしたい方に'],
            ['スタンダード', 'シンプル＋部屋の明渡し・遺品整理・通知', '33万円〜', '賃貸にお住まいの方におすすめ'],
            ['フルサポート', 'スタンダード＋デジタル整理・遺言執行', '44万円〜', '全てをおまかせしたい方に'],
        ],
        col_widths=[1.2, 2.3, 1.3, 1.8]
    )

    add_styled_para(doc, '※上記は契約時の費用です。実際にお手伝いする際の報酬（執行時報酬）は別途かかります。', size=Pt(9), color=MED_GRAY)
    add_styled_para(doc, '※「遺産清算方式」をお選びいただくと、預託金なしでご契約いただけます。詳しくはご相談時にご説明します。', size=Pt(9), color=MED_GRAY, space_after=Pt(12))

    add_divider(doc, '2E75B6')

    # よくある質問
    p = doc.add_paragraph()
    run = p.add_run('よくあるご質問')
    set_font(run, size=Pt(14), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(8)

    qas = [
        ('Q. 弁護士に頼むと高いのでは？', 'A. 契約時の費用は22万円からです。一般的なNPOや他の事業者と比べても、同等かそれ以下の水準です。弁護士だからこそ、法的なトラブルにも対応できる安心感があります。'),
        ('Q. 契約した後、気が変わったらどうなりますか？', 'A. いつでも解約可能です。解約時に違約金等は発生しません。'),
        ('Q. 遠方に住んでいても大丈夫ですか？', 'A. お電話やオンラインでのご相談も承っております。お気軽にお問い合わせください。'),
        ('Q. 身寄りがまったくいなくても利用できますか？', 'A. はい。身寄りのない方にこそ、ご利用いただきたいサービスです。弁護士が全てを責任を持ってお引き受けします。'),
    ]
    for q, a in qas:
        p = doc.add_paragraph()
        run = p.add_run(q)
        set_font(run, size=Pt(10.5), color=DARK_BLUE, bold=True)
        p.paragraph_format.space_after = Pt(2)

        p = doc.add_paragraph()
        run = p.add_run(a)
        set_font(run, size=Pt(10), color=DARK_GRAY)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Inches(0.2)

    doc.add_page_break()

    # ===== 4ページ目：CTA =====
    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('まずはお気軽に\nお電話ください')
    set_font(run, size=Pt(20), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.4

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ご相談は無料です')
    set_font(run, size=Pt(12), color=MID_BLUE)
    p.paragraph_format.space_after = Pt(16)

    add_divider(doc, 'B8860B')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TEL)
    set_font(run, size=Pt(32), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'受付時間：{HOURS}')
    set_font(run, size=Pt(11), color=MED_GRAY)
    p.paragraph_format.space_after = Pt(24)

    add_divider(doc, 'B8860B')

    doc.add_paragraph()

    # 事務所情報
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, size, bold, color in [
        (OFFICE_NAME, Pt(13), True, NAVY),
        ('\n', Pt(4), False, DARK_GRAY),
        (REPRESENTATIVE, Pt(11), False, DARK_GRAY),
        ('\n', Pt(6), False, DARK_GRAY),
        (ADDRESS, Pt(9.5), False, MED_GRAY),
        ('\n', Pt(2), False, DARK_GRAY),
        (URL, Pt(9.5), False, MID_BLUE),
        ('\n', Pt(8), False, DARK_GRAY),
        ('東京弁護士会所属', Pt(9), False, MED_GRAY),
        ('\n', Pt(2), False, DARK_GRAY),
        ('初回相談無料', Pt(11), True, DARK_BLUE),
    ]:
        run = p.add_run(text)
        set_font(run, size=size, color=color, bold=bold)

    doc.save(r'C:\Users\user\Desktop\claudeマスター\【リーフレット】おひとりさまの安心サポート.docx')
    print('✓ 汎用リーフレット 生成完了')


# ===== 実行 =====
if __name__ == '__main__':
    create_dm_for_management()
    create_newspaper_ad()
    create_leaflet()
    print('\n全3ドキュメント生成完了！')
