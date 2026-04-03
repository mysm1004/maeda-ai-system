#!/usr/bin/env python3
"""法務AI活用支援事業 企画書 DOCX生成スクリプト"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ===== スタイル設定 =====
style = doc.styles['Normal']
font = style.font
font.name = 'Yu Gothic'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Yu Gothic'
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    if level == 1:
        hs.font.size = Pt(18)
        hs.font.bold = True
    elif level == 2:
        hs.font.size = Pt(14)
        hs.font.bold = True
    else:
        hs.font.size = Pt(12)
        hs.font.bold = True

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        set_cell_shading(cell, '1a1a2e')
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, 'f0f0f5')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)
    return p

# ============================================================
# 表紙
# ============================================================
for _ in range(8):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('法務AI活用支援事業')
r.font.size = Pt(32)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('企 画 書')
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('弁護士法人 東京新橋法律事務所')
r.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('代表弁護士 前田 祥夢')
r.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('2026年4月')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============================================================
# 目次
# ============================================================
doc.add_heading('目次', level=1)
toc_items = [
    '1. エグゼクティブサマリー',
    '2. 市場環境',
    '3. 競合分析',
    '4. サービス設計',
    '5. ターゲット',
    '6. 販売・集客戦略',
    '7. 収支計画',
    '8. リスクと対策',
    '9. 助成金・補助金の活用',
    '10. 今後のアクション',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(6)
    p.runs[0].font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 1. エグゼクティブサマリー
# ============================================================
doc.add_heading('1. エグゼクティブサマリー', level=1)

doc.add_heading('事業コンセプト', level=2)
p = doc.add_paragraph()
r = p.add_run('法務部が「安心して、効果を測って、自社の資産として」AIを使える状態を作る、弁護士による法務AI活用支援事業。')
r.font.size = Pt(11)

doc.add_heading('差別化の3本柱', level=2)
add_bullet(doc, 'ハルシネーション統制 — ', '① ')
p = doc.add_paragraph('弁護士が品質基準を設計し、法務AIの「もっともらしい嘘」を仕組みとして防ぐ')
p.paragraph_format.left_indent = Cm(1.5)
add_bullet(doc, '運用基準策定 — ', '② ')
p = doc.add_paragraph('自社完結と弁護士相談の線引きを明確化。グリーン/イエロー/レッドの3段階で運用基準書を策定・納品')
p.paragraph_format.left_indent = Cm(1.5)
add_bullet(doc, '定量ROI管理 — ', '③ ')
p = doc.add_paragraph('リクルート式KPIで導入効果を数字で証明。ベースライン計測→月次ダッシュボード→経営報告まで一気通貫')
p.paragraph_format.left_indent = Cm(1.5)

doc.add_heading('初年度目標', level=2)
add_bullet(doc, '売上: 3,000万円 / 粗利: 2,600万円')
add_bullet(doc, 'フロント商品受注: 8〜10件')
add_bullet(doc, '月額顧問ストック: 4〜6社')
add_bullet(doc, '法務AIエージェント構築: 2件')

doc.add_page_break()

# ============================================================
# 2. 市場環境
# ============================================================
doc.add_heading('2. 市場環境', level=1)

doc.add_heading('市場規模', level=2)
add_bullet(doc, '日本のリーガルテック市場: 電子契約だけで2025年約395億円、CAGR 11.3%')
add_bullet(doc, 'グローバル法務AI市場: 2025年31.1億ドル → 2030年108億ドル（CAGR 28.3%）')

doc.add_heading('企業の課題（データ裏付け）', level=2)
add_bullet(doc, '52.9%の企業が「AI導入の効果はまだ不明」（DirectCloud調査）')
add_bullet(doc, '約8割が生成AIを導入しているが、同じ割合が「収益面の成果なし」（マッキンゼー）')
add_bullet(doc, '約7割の法務担当者が契約書の見落としを経験（LegalOn調査）')
add_bullet(doc, '「ツールは入れたが成果が出ない」企業が多数')

doc.add_heading('空白地帯', level=2)
p = doc.add_paragraph()
r = p.add_run('「法務AIを"事故なく・基準を持って・数字で効果を証明しながら"運用できる状態を、弁護士が一緒に作る」プレイヤーは市場にゼロ。')
r.font.size = Pt(11)
r.bold = True

doc.add_page_break()

# ============================================================
# 3. 競合分析
# ============================================================
doc.add_heading('3. 競合分析', level=1)

doc.add_heading('競合カテゴリと限界', level=2)
add_table(doc,
    ['カテゴリ', '代表プレイヤー', '売っているもの', '本質的な限界'],
    [
        ['SaaSツール', 'LegalOn, MNTSQ,\nGVA OLGA', '高精度なAIツール', 'ツールの外は管轄外。使いこなし・検証・運用基準は顧客任せ'],
        ['法務BPO', 'Authense\n法務クラウド', '弁護士リソース', 'AI活用の設計がない。人を増やすモデルでAI時代と逆行'],
        ['大手コンサル', 'PwC, デロイト', 'AI戦略の絵', '絵を描いて終わり。実装しない。法務の現場を知らない'],
        ['AI専業コンサル', 'AI総研, AVILEN等', '汎用AI導入の技術', '法務の専門性ゼロ。ハルシネーションの法的リスクを判断できない'],
        ['AI研修会社', 'SHIFT AI等', 'AI活用の知識', '教えて終わり。定着・検証・効果測定なし'],
        ['品質保証', 'SHIFT', 'AI出力の品質テスト', '法的判断の正誤は評価不能'],
    ],
    col_widths=[3.5, 3.5, 4, 6]
)

doc.add_paragraph()
doc.add_heading('当社のポジション', level=2)
p = doc.add_paragraph()
r = p.add_run('「弁護士資格を持ったAI実装者」という唯一のポジション。')
r.bold = True
r.font.size = Pt(11)
p.add_run('弁護士法72条の壁、ビジネスモデルの矛盾、ケイパビリティの不足のいずれかが競合の参入障壁になり、構造的に真似できない。')

doc.add_paragraph()
doc.add_heading('サービス要素別比較（3つの急所）', level=2)
add_table(doc,
    ['サービス要素', '競合の状況', '当社'],
    [
        ['ハルシネーション統制\n（検証プロセス設計）', '提供プレイヤーゼロ', '弁護士が法務特化で構築'],
        ['運用基準策定\n（自社完結/弁護士の線引き）', '提供プレイヤーゼロ', '運用基準書を弁護士が策定・納品'],
        ['定量ROI管理\n（ベースライン〜月次KPI）', '提供プレイヤーゼロ', 'リクルート式KPIで月次レポート'],
    ],
    col_widths=[5, 4, 5.5]
)

doc.add_page_break()

# ============================================================
# 4. サービス設計
# ============================================================
doc.add_heading('4. サービス設計', level=1)

doc.add_heading('商品ラインナップ（4層構造）', level=2)

# 商品S
doc.add_heading('商品S: 無料セミナー＆個別相談', level=3)
p = doc.add_paragraph('リード獲得。月1回ウェビナー開催。個別相談30分（無料）で商談化。')

# 商品A
doc.add_heading('商品A: 法務AI研修（50万円/単発）', level=3)
add_bullet(doc, '法務部員向け実践研修 12時間（半日×4回）')
add_bullet(doc, 'Day1: 生成AI基礎＋法務での活用方法＋プロンプト実践')
add_bullet(doc, 'Day2: ハルシネーション対策＋運用ルール＋ワークショップ')
add_bullet(doc, '人材開発支援助成金（リスキリング支援コース）対応可')
add_bullet(doc, '前田工数: 20h / 粗利: 50万円（粗利率100%）')

# 商品B
doc.add_heading('商品B: 法務AI導入プログラム（150万円/3ヶ月）', level=3)
p = doc.add_paragraph()
r = p.add_run('Phase 1: 診断・設計（1ヶ月目）— 工数15h')
r.bold = True
add_bullet(doc, '法務業務フロー診断＋ヒアリング')
add_bullet(doc, 'ベースラインKPI計測の設計・実施')
add_bullet(doc, 'AI化対象領域の特定＋ROI試算')
add_bullet(doc, '運用基準書（グリーン/イエロー/レッド）のドラフト')

p = doc.add_paragraph()
r = p.add_run('Phase 2: 実装（2ヶ月目）— 工数20h')
r.bold = True
add_bullet(doc, 'プロンプトテンプレート設計')
add_bullet(doc, 'ハルシネーション検証プロセス構築')
add_bullet(doc, 'ツール設定支援（ChatGPT Enterprise等）')
add_bullet(doc, '法務部員向けハンズオン研修 半日×2回')

p = doc.add_paragraph()
r = p.add_run('Phase 3: 定着・計測（3ヶ月目）— 工数10h')
r.bold = True
add_bullet(doc, '運用基準書のFix＋納品')
add_bullet(doc, 'Slack/Teamsでの質問対応')
add_bullet(doc, '導入効果レポート（ベースライン比較）作成')

p = doc.add_paragraph()
r = p.add_run('納品物: ')
r.bold = True
p.add_run('法務AI運用基準書、プロンプトテンプレート集（20〜30本）、ハルシネーション検証フローマニュアル、導入効果レポート、研修資料')
p = doc.add_paragraph()
r = p.add_run('前田工数: 45h / 粗利: 150万円（粗利率100%）')
r.bold = True

# 商品C
doc.add_heading('商品C: 法務AIエージェント構築（300〜500万円/3〜6ヶ月）', level=3)
add_bullet(doc, '契約レビューエージェント — 自社過去契約書・審査基準をRAGに格納。リスク自動分類。', '① ')
add_bullet(doc, '法務QAエージェント — 社内規程・過去QAをナレッジベース化。自信度スコア付き自動回答。', '② ')
add_bullet(doc, '法令調査エージェント — 法改正情報の自動収集＋要約＋影響度評価。', '③ ')
add_bullet(doc, '運用基盤 — ハルシネーション統制レイヤー、利用ログ・KPIダッシュボード。', '④ ')
p = doc.add_paragraph()
r = p.add_run('技術スタック: ')
r.bold = True
p.add_run('Claude API or Azure OpenAI / Dify or カスタムRAG / n8n or カスタム / Slack・Teams連携')
p = doc.add_paragraph()
r = p.add_run('前田工数: 60〜80h / エンジニア外注: 150〜250万円 / 粗利: 150〜250万円（粗利率50〜55%）')
r.bold = True

# 商品D
doc.add_heading('商品D: 法務AI運用顧問（月額ストック）', level=3)
add_table(doc,
    ['プラン', '月額', '内容', '前田工数', '粗利'],
    [
        ['スタンダード', '20万円', 'AI運用相談（無制限）/プロンプト月次更新/\nハルシネーションサンプルレビュー(5件)/\n月次KPIレポート/法務相談(月3h)', '月5〜6h', '月20万'],
        ['プレミアム', '35万円', '上記全て+月1回定例/新規AI活用検討(四半期1テーマ)/\n法務相談(月8h)/四半期経営報告', '月10〜12h', '月35万'],
    ],
    col_widths=[3, 2, 6, 2.5, 2]
)

doc.add_paragraph()
doc.add_heading('粗利サマリー', level=2)
add_table(doc,
    ['商品', '価格', '前田工数', '外注費', '粗利', '粗利率', '時間単価'],
    [
        ['A: 研修', '50万', '20h', '0', '50万', '100%', '2.5万/h'],
        ['B: 導入PG', '150万', '45h', '0', '150万', '100%', '3.3万/h'],
        ['C: エージェント構築', '300〜500万', '60〜80h', '150〜250万', '150〜250万', '50〜55%', '2.5〜3.4万/h'],
        ['D: 月額顧問(S)', '月20万', '月6h', '0', '月20万', '100%', '3.3万/h'],
        ['D: 月額顧問(P)', '月35万', '月12h', '0', '月35万', '100%', '2.9万/h'],
    ],
    col_widths=[3.5, 2.5, 2, 2.5, 2.5, 1.5, 2]
)

doc.add_page_break()

# ============================================================
# 5. ターゲット
# ============================================================
doc.add_heading('5. ターゲット', level=1)

doc.add_heading('メインターゲット', level=2)
add_bullet(doc, '従業員500〜5,000人の中堅企業')
add_bullet(doc, '法務部5〜20人')
add_bullet(doc, 'AIツールを導入済み or 検討中だが成果が出ていない企業')

doc.add_heading('ターゲット選定理由', level=2)
add_table(doc,
    ['セグメント', '法務部規模', '特徴', '評価'],
    [
        ['大企業', '50人以上', '既にLegalOn等を導入済み。意思決定が遅い', '△'],
        ['中堅企業', '5〜20人', '最有望。ツールを入れたが使いこなせていない。予算もある。意思決定が速い', '◎'],
        ['スタートアップ', '1〜2人', '予算が厳しいが経営者判断で即決あり', '○'],
    ],
    col_widths=[3, 3, 7.5, 2]
)

doc.add_page_break()

# ============================================================
# 6. 販売・集客戦略
# ============================================================
doc.add_heading('6. 販売・集客戦略', level=1)

doc.add_heading('セールスファネル', level=2)
add_table(doc,
    ['ステップ', '目標', '施策'],
    [
        ['① 認知', '月間リーチ\n3,000〜5,000人', '法務メディア寄稿（BUSINESS LAWYERS等/会員12万人）\nウェビナー月1回（50〜100名）\nLinkedIn/X/note発信'],
        ['② リード獲得', '月間50〜100件', 'セミナー参加/ホワイトペーパーDL\nHP3本: AI活用度チェックリスト/運用基準テンプレ/ROI計算シート'],
        ['③ 無料診断', '月間10〜15件', '30分オンライン「法務AI活用度チェック」\n5段階スコア＋ROI試算を無料で提供'],
        ['④ フロント受注', '月間1.5〜2件', 'ROIで投資の正当性を証明\n月50万×3回払いで部長決裁圏内に'],
        ['⑤ バックエンド移行', '受注の50〜70%', 'フロント完了時の導入効果レポートで自然移行'],
    ],
    col_widths=[3, 3.5, 9]
)

doc.add_paragraph()
doc.add_heading('初動アクションプラン', level=2)
add_table(doc,
    ['時期', 'アクション'],
    [
        ['Week 1', 'リクルート人脈に連絡。法務部長と繋がる紹介を依頼'],
        ['Week 2', '既存顧問先にAI活用支援の追加提案'],
        ['Week 3', 'LinkedIn発信開始（弁護士×AI活用の実践ネタ）'],
        ['Week 4', 'BUSINESS LAWYERSに寄稿打診メール'],
    ],
    col_widths=[3, 13]
)

doc.add_paragraph()
doc.add_heading('3フェーズの時系列計画', level=2)
add_table(doc,
    ['フェーズ', '期間', '内容', '受注目標'],
    [
        ['A: 仕込み期', '1〜3ヶ月目', 'コンテンツ制作、SNS発信開始、\n紹介営業', '1件'],
        ['B: 立ち上げ期', '4〜6ヶ月目', 'ウェビナー開始、寄稿公開、\nリード育成', '月1〜2件'],
        ['C: 成長期', '7〜12ヶ月目', 'ウェビナー月2回、事例蓄積、\n紹介の連鎖', '月2件+\nストック10社'],
    ],
    col_widths=[3, 3, 6, 3.5]
)

doc.add_page_break()

# ============================================================
# 7. 収支計画
# ============================================================
doc.add_heading('7. 収支計画', level=1)

doc.add_heading('初年度（保守的シミュレーション）', level=2)
add_table(doc,
    ['項目', '件数', '売上', '工数'],
    [
        ['商品B: 導入PG', '8件', '1,200万円', '360h'],
        ['商品A: 研修', '4件', '200万円', '80h'],
        ['商品C: エージェント構築', '2件', '800万円', '150h'],
        ['商品D: 月額顧問', '平均4社', '800万円', '280h'],
        ['売上合計', '—', '3,000万円', '870h'],
        ['マーケ投資', '—', '▲400万円', '—'],
        ['エンジニア外注', '—', '▲400万円', '—'],
        ['粗利', '—', '2,200万円', '—'],
        ['月平均稼働', '—', '—', '73h/月'],
    ],
    col_widths=[5, 3, 4, 3.5]
)

doc.add_paragraph()
doc.add_heading('2年目（成長期）', level=2)
add_table(doc,
    ['項目', '売上'],
    [
        ['商品B', '1,800万円'],
        ['商品A', '300万円'],
        ['商品C', '1,800万円'],
        ['商品D', '3,240万円'],
        ['売上合計', '7,140万円'],
        ['外注・マーケ費', '▲1,400万円'],
        ['アシスタント人件費', '▲480万円'],
        ['粗利', '5,260万円'],
    ],
    col_widths=[8, 8]
)

doc.add_page_break()

# ============================================================
# 8. リスクと対策
# ============================================================
doc.add_heading('8. リスクと対策', level=1)
add_table(doc,
    ['リスク', '対策'],
    [
        ['LegalOn等が導入コンサルに参入', '「ツール非依存」のポジションを堅持。複数ツール横断の設計力が差別化'],
        ['AI進化で導入支援ノウハウが陳腐化', 'エージェント構築（商品C）に重心を移す。AIが進化するほど構築需要は増える'],
        ['スケーラビリティの壁', '2年目にアシスタント採用＋エンジニアパートナーの固定化'],
        ['「弁護士がAI？」の信頼性ギャップ', '自事務所のAI活用実績をデモで見せる。ウェビナーでライブ実演'],
        ['法務部の保守性（意思決定の遅さ）', '無料診断でROI数値を見せて稟議資料まで用意。月50万×3回払いで部長決裁圏内に'],
    ],
    col_widths=[6, 10.5]
)

doc.add_page_break()

# ============================================================
# 9. 助成金・補助金
# ============================================================
doc.add_heading('9. 助成金・補助金の活用', level=1)
add_table(doc,
    ['制度', '適用可否', '備考'],
    [
        ['人材開発支援助成金\n（リスキリング支援コース）', '商品Aのみ\n適用可', '10h以上のOFF-JT。中小企業75%助成。経費上限30万円'],
        ['デジタル化・\nAI導入補助金', '顧客が中小企業\nなら適用可', 'IT導入支援事業者登録が必要。補助率1/2〜4/5'],
        ['東京都\nDX推進助成金', '都内中小企業\nなら有力', 'DXコンサル費対象。助成率1/2、上限300万円'],
    ],
    col_widths=[4.5, 4, 8]
)
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('※ 助成金はオプション扱い。メインの売り方はROI訴求。')
r.italic = True
r.font.size = Pt(10)

doc.add_page_break()

# ============================================================
# 10. 今後のアクション
# ============================================================
doc.add_heading('10. 今後のアクション', level=1)

actions = [
    'サービスLP制作（3パターン）',
    'ホワイトペーパー3本の制作（AI活用度チェックリスト/運用基準テンプレート/ROI計算シート）',
    'ウェビナー資料の制作',
    'BUSINESS LAWYERS寄稿記事の執筆（3本シリーズ）',
    'エンジニアパートナーの選定',
    '紹介営業の開始（リクルート人脈＋既存顧問先）',
]
for i, action in enumerate(actions, 1):
    p = doc.add_paragraph(f'{i}. {action}')
    p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('以上')
r.font.size = Pt(12)

# ===== 保存 =====
out_path = r'C:\Users\user\Desktop\claudeマスター\案件\法務AI事業\企画書_法務AI活用支援事業.docx'
doc.save(out_path)
print('Done: ' + out_path)
