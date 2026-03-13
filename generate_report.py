#!/usr/bin/env python3
"""死後事務委任 検討レポート生成スクリプト"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import datetime

doc = Document()

# ===== スタイル設定 =====
style = doc.styles['Normal']
font = style.font
font.name = 'Yu Gothic'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Yu Gothic'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    if level == 1:
        h.font.size = Pt(18)
        h.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)
    elif level == 2:
        h.font.size = Pt(14)
        h.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    elif level == 3:
        h.font.size = Pt(12)
        h.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

# ページ設定（A4）
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)


def add_para(text, bold=False, size=None, color=None, space_after=Pt(6), alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Yu Gothic'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    if bold:
        run.bold = True
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = space_after
    if alignment:
        p.alignment = alignment
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    for run in p.runs:
        run.font.name = 'Yu Gothic'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
        run.font.size = Pt(10.5)
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ヘッダー
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.name = 'Yu Gothic'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '2E75B6',
            qn('w:val'): 'clear'
        })
        shading.append(shading_elm)

    # データ行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Yu Gothic'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
            if r_idx % 2 == 1:
                shading = cell._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F2F7FB',
                    qn('w:val'): 'clear'
                })
                shading.append(shading_elm)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Yu Gothic'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True
    return p


# ===== 表紙 =====
doc.add_paragraph()
doc.add_paragraph()
add_para('前田法律事務所', size=Pt(14), color=RGBColor(0x66, 0x66, 0x66), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('死後事務委任契約 特化サービス', size=Pt(24), bold=True, color=RGBColor(0x1A, 0x47, 0x7A), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('事業構想 検討レポート', size=Pt(18), color=RGBColor(0x2E, 0x75, 0xB6), alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para(f'作成日：2026年3月12日', size=Pt(11), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Phase 1（壁打ち）・Phase 2（戦略設計）統合版', size=Pt(11), alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para('本レポートは、AI壁打ちシステムによる多角的検討の結果をまとめたものです。', size=Pt(10), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para('競合分析、法的リスク、ターゲット行動分析、価格設計、集客戦略を網羅しています。', size=Pt(10), alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ===== 目次 =====
doc.add_heading('目次', level=1)
toc_items = [
    'Phase 1：壁打ち（事業構想の検討）',
    '  1. 事業概要',
    '  2. 市場環境',
    '  3. 競合分析',
    '  4. 法的リスクの検討',
    '  5. ターゲット分析の深掘り',
    '  6. 提携先の検討と優先度',
    '  7. 独り身高齢者の日常生活実態',
    '',
    'Phase 2：戦略設計',
    '  1. 預託金の管理方法',
    '  2. 弁護士費用の値付け',
    '  3. 集客チャネルとCPA分析',
    '  4. メッセージング戦略',
    '  5. 最終チャネル設計',
    '  6. 段階別ロードマップ',
    '  7. 課題・ネック・反対意見の整理',
]
for item in toc_items:
    if item == '':
        doc.add_paragraph()
    elif item.startswith('Phase'):
        add_para(item, bold=True, size=Pt(11))
    else:
        add_para(item, size=Pt(10.5))

doc.add_page_break()

# =====================================================
# Phase 1
# =====================================================
doc.add_heading('Phase 1：壁打ち（事業構想の検討）', level=1)

# --- 1. 事業概要 ---
doc.add_heading('1. 事業概要', level=2)
add_para('前田法律事務所が新たに展開を検討する「死後事務委任契約」特化サービスの事業構想について、多角的な検討を行った。')

doc.add_heading('サービスの定義', level=3)
add_bullet('弁護士が生前に依頼者と死後事務委任契約を締結')
add_bullet('依頼者の死亡後、弁護士が以下を代行：死亡届提出、葬儀手配、納骨、行政手続き、賃貸明渡し、遺品整理、デジタルアカウント削除等')
add_bullet('身元保証・任意後見は提供しない（死後事務委任に特化することで差別化）')

doc.add_heading('ターゲット', level=3)
add_bullet('独り身高齢者（配偶者なし・子なし、または家族と疎遠な高齢者）')
add_bullet('自分の死後の手続きを誰にも頼めない層')

doc.add_heading('特化戦略の意図', level=3)
add_bullet('身元保証・任意後見まで手を広げると、きずなの会等の大手NPOと正面衝突する')
add_bullet('死後事務委任に絞ることで、弁護士ならではの法的専門性（遺言執行、紛争対応）が活きる')
add_bullet('サービス範囲を限定することで、料金の明瞭化・低価格化が可能')

# --- 2. 市場環境 ---
doc.add_heading('2. 市場環境', level=2)

doc.add_heading('市場規模', level=3)
add_table(
    ['指標', '数値', '出典'],
    [
        ['65歳以上の単身世帯', '855万世帯（2023年）', '国勢調査'],
        ['2050年の独居高齢者予測', '1,084万人', '国立社会保障・人口問題研究所'],
        ['年間孤独死数', '約6.8万人（2024年）', '警察庁'],
        ['おひとりさまの終活実践率', 'わずか13%', '燦HD調査（2023年）'],
        ['高齢単身世帯の借家率', '32.2%', '住宅・土地統計調査'],
        ['2050年の男性単身世帯の未婚率', '59.7%', '日本総研'],
    ],
    col_widths=[2.5, 2.0, 2.0]
)
add_note('需要は急拡大している一方、供給側（特に弁護士による死後事務委任サービス）は未成熟。')

# --- 3. 競合分析 ---
doc.add_heading('3. 競合分析', level=2)

doc.add_heading('Tier1：大手・全国展開', level=3)
add_table(
    ['競合', '料金目安', '葬儀手配', '弱点'],
    [
        ['りすシステム', '230〜330万円', '○', '高額すぎる、預託金不透明'],
        ['きずなの会', '130〜230万円', '○', '名古屋中心、全国カバー薄い'],
        ['イオンライフ', '80〜150万円', '◎', '死後事務の法的専門性が弱い（弁護士不在）'],
        ['小さなお葬式', '9.9〜70万円', '◎', '死後事務委任は未対応'],
        ['終活協議会（心託）', '68〜118万円', '△', '弁護士ではない'],
    ],
    col_widths=[1.5, 1.5, 1.0, 2.5]
)

doc.add_heading('最大の競合：きずなの会 × 名城法律事務所', level=3)
add_para('「弁護士 × 死後事務委任」の先行モデルが既に存在する。', bold=True)

add_table(
    ['項目', '内容'],
    [
        ['関係性', '名城法律事務所の弁護士（西山富夫氏）がきずなの会を設立。事実上の一体組織'],
        ['累計会員数', '15,600名超（年間800名ペースで増加中）'],
        ['全国拠点', '16か所（東海・関東・関西中心）'],
        ['弁護士の役割', '金銭預託管理・後見・法律相談（裏方）'],
        ['NPOの役割', '身元保証・生活支援・葬送支援（表の顔）'],
        ['推定収益', '金銭預託手数料だけで年約2億円のストック収入'],
    ],
    col_widths=[1.8, 4.7]
)

add_para('きずなの会の弱点（差別化のポイント）：', bold=True)
add_bullet('料金が不透明・高額（初期数十万＋預託金で100万円超）')
add_bullet('NPOスタッフが対応し、弁護士は裏方（パーソナル感の欠如）')
add_bullet('東海・関東・関西中心で地方は空白地帯')
add_bullet('デジタル終活（SNS削除・暗号資産等）は未対応')

add_note('結論：完全なブルーオーシャンではないが、差別化の余地は十分にある。')

# --- 4. 法的リスク ---
doc.add_heading('4. 法的リスクの検討', level=2)
add_para('検討過程で浮上した法的リスクを網羅的に整理する。いずれも事業設計の根幹に関わる重要論点。', bold=True)

doc.add_heading('4-1. 個人情報保護法（葬儀社・寺院のリスト共有）', level=3)
add_para('当初の構想では「葬儀社や寺院の顧客リストにDMを送付する」ことを検討していたが、以下の理由で原則違法と判断した。')

add_table(
    ['方法', '適法性', '問題点'],
    [
        ['そのまま第三者提供', '違法（法27条1項違反）', '本人同意なき第三者提供'],
        ['オプトアウト届出', '理論上可能だが非現実的', '葬儀社が「名簿販売業者」扱いに。宗教情報は要配慮個人情報で不可'],
        ['共同利用（法27条5項3号）', '条件付きで可能', '取得時に共同利用の旨を明示していない既存リストには遡及適用が困難'],
        ['委託構成', '該当しない', '弁護士が自己の顧客獲得目的で送るため委託にならない'],
    ],
    col_widths=[1.8, 1.8, 3.0]
)

doc.add_heading('4-2. 寺院の檀家名簿：刑法上のリスク', level=3)
add_bullet('寺院の宗教活動は個人情報保護法の適用除外（法57条1項3号）')
add_bullet('しかし、弁護士のDM目的は宗教活動ではないため除外に該当しない')
add_bullet('刑法134条2項（秘密漏示罪）：宗教者が業務上知り得た秘密を正当な理由なく漏らすと6か月以下の懲役または10万円以下の罰金')
add_para('結論：寺院の檀家名簿を営業目的で外部に渡すのは刑事罰リスクあり。', bold=True)

doc.add_heading('4-3. 弁護士広告規制（日弁連規程）', level=3)
add_bullet('規程第6条：特定の事件の当事者で面識のない者への勧誘DMは原則禁止')
add_bullet('葬儀社リストから遺族にDMを送る場合、相続事件の当事者への勧誘に該当する可能性')
add_bullet('違反時は戒告・業務停止（最長2年）・退会命令の対象')

doc.add_heading('4-4. 弁護士法72条（紹介料の禁止）', level=3)
add_bullet('非弁護士が報酬を得て法律事件の周旋を行うことは禁止（弁護士法72条）')
add_bullet('弁護士職務基本規程13条：事件の紹介を受けた弁護士は紹介料を支払ってはならない')
add_bullet('賃貸管理会社やサ高住に「1件○万円」の紹介料を支払うことは違法')
add_bullet('合法な代替策：セミナー共催、勉強会講師、契約書ひな形の無償提供')

add_note('法的リスクの検討により、「リスト共有によるDM送付」は全面的に撤回。提携先との関係構築は「課題解決ベース」で行うこととした。')

# --- 5. ターゲット分析 ---
doc.add_heading('5. ターゲット分析の深掘り', level=2)

doc.add_heading('5-1. 「葬儀社のリストでDM」の前提崩壊', level=3)
add_para('ターゲットである「独り身高齢者」について深掘りした結果、当初の集客前提が根本から崩れることが判明した。')

add_bullet('互助会は「家族前提の制度」であり、単身者にはメリットが薄い')
add_bullet('互助会に加入していても、死亡時にサービスを発動する人がいないという致命的欠陥')
add_bullet('葬儀社の主なリードは「親の葬儀を考える50〜60代の子供」から。独り身高齢者本人からではない')
add_bullet('葬儀社は「家族構成」でセグメントしたリストを通常持っていない')
add_bullet('おひとりさまの終活実践率はわずか13%（燦HD調査）')

add_para('戦略転換の方向性：', bold=True)
add_para('「葬儀社 → 高齢者」ではなく「高齢者 → 弁護士 → 葬儀社」の流れに逆転。葬儀社はリスト提供元ではなく、受任後の「出口パートナー」として位置づける。')

# --- 6. 提携先 ---
doc.add_heading('6. 提携先の検討と優先度', level=2)
add_para('「独り身高齢者との接点」「提携先のメリットの大きさ」「提携のしやすさ」「スケーラビリティ」の4軸で評価。')

doc.add_heading('提携先メリットランキング', level=3)
add_table(
    ['順位', '提携先', 'メリットの種類', '金額インパクト', '切実度'],
    [
        ['1', '賃貸管理会社', '損失回避（孤独死1件=60万+家賃下落）', '数百万の損失防止', '極めて切実'],
        ['2', '葬儀社', '売上直結（葬儀1件=30〜100万の確約予約）', '直接売上', '切実'],
        ['3', 'サ高住', '入居率向上（身寄りなし受入可能に）', '年間数百万', '切実'],
        ['4', '家賃保証会社', '審査通過率UP→契約数増', '間接的売上', 'やや切実'],
        ['5', 'ケアマネ', '業務リスク軽減（無権限行為の回避）', '金銭的ではない', '低い'],
        ['6', '社協・地域包括', '紹介先があることで業務完結', '金銭的ではない', '低い'],
    ],
    col_widths=[0.5, 1.3, 2.0, 1.5, 1.2]
)

doc.add_heading('Tier 1：最優先提携先', level=3)

add_para('1位：賃貸管理会社', bold=True, size=Pt(11))
add_bullet('孤独死発生時の損害：原状回復平均60万円＋事故物件化で家賃20〜30%下落')
add_bullet('国交省＋法務省がモデル契約条項で死後事務委任を公式推奨（2021年）')
add_bullet('営業トーク：「大家の孤独死リスクをゼロにする」')
add_bullet('導線：管理会社が高齢入居者に「死後事務委任契約」を入居条件として推奨')

add_para('2位：葬儀社（出口パートナーとして）', bold=True, size=Pt(11))
add_bullet('リスト目的ではなく「確約予約を流す」パートナー')
add_bullet('おひとりさまの事前予約は「死亡を連絡する人がいない」「費用を払う人もいない」問題がある')
add_bullet('死後事務委任の弁護士が介在すれば、死亡通知→葬儀発注→費用支払い→遺骨受取が全て確約される')
add_bullet('弁護士は葬儀社にとって「最も確実な送客チャネル」になる')

add_para('3位：サ高住', bold=True, size=Pt(11))
add_bullet('全国約28万戸。入居者の多くは自立〜要介護1で意思能力あり')
add_bullet('退去理由の40.3%が「死亡による契約終了」')
add_bullet('入居時の契約フローに死後事務委任をセットで組み込める')

doc.add_heading('施設種別と意思能力の問題', level=3)
add_table(
    ['施設種別', '入居者の状態', '意思能力', '契約可否'],
    [
        ['サ高住', '自立〜要介護1が中心', '◎ あり', '契約可能'],
        ['住宅型有料老人ホーム', '自立〜要介護3', '○ 多くはあり', '入居時に案内可能'],
        ['介護付き有料老人ホーム', '要介護1〜5', '△ ケースバイケース', '早期入居者なら可能'],
        ['特養', '原則要介護3以上', '✕ 多くは困難', '認知症進行で難しい'],
        ['グループホーム', '認知症が入居要件', '✕', '契約能力なし'],
    ],
    col_widths=[2.0, 1.8, 1.5, 1.5]
)
add_note('結論：サ高住が最も相性がいい。特養・グループホームは意思能力の問題で対象外。')

doc.add_heading('社会福祉協議会と地域包括支援センターの違い', level=3)
add_table(
    ['', '社会福祉協議会（社協）', '地域包括支援センター'],
    [
        ['一言で', '地域福祉の総合司令塔', '高齢者のよろず相談窓口'],
        ['法的根拠', '社会福祉法', '介護保険法'],
        ['設置単位', '全市区町村に1つ', '中学校区に1つ（全国5,400か所超）'],
        ['運営主体', '社会福祉法人（半官半民）', '自治体が委託（社協が受託することも多い）'],
        ['接点の深さ', '金銭管理・見守り事業で深い関係', 'ケアプラン作成・介護相談で日常的に接触'],
    ],
    col_widths=[1.2, 2.7, 2.7]
)

# --- 7. 日常生活実態 ---
doc.add_heading('7. 独り身高齢者の日常生活実態', level=2)
add_para('集客チャネルを設計するにあたり、ターゲットである独り身高齢者の日常を徹底調査した。「彼らが普段何を見聞きし、何に触れ、何を見て過ごしているか」を把握することが最重要。')

doc.add_heading('1日のタイムスケジュール（データから再構成）', level=3)
add_table(
    ['時間', '行動', '接触メディア'],
    [
        ['6:00', '起床。テレビをつける（NHK）', 'テレビ'],
        ['6:30', '新聞を読む（朝刊）', '新聞'],
        ['7:00', '朝食（自炊 or パン）', ''],
        ['8:00〜', 'テレビ or ラジオ。5時間以上つけっぱなしの人が4割', 'テレビ・ラジオ'],
        ['9:00', '通院（月2回以上が6割）or 買い物（スーパー。週3〜4回）', '外出先の掲示物'],
        ['12:00', '昼食（一人で）', ''],
        ['13:00', 'テレビ。昼のワイドショー', 'テレビ'],
        ['14:00', '散歩 or 特にやることがない', ''],
        ['19:00', 'テレビ（NHKニュース7）', 'テレビ'],
        ['22:00', '就寝', ''],
    ],
    col_widths=[1.0, 3.5, 1.5]
)
add_note('70代以上の1日のテレビ視聴：平均5時間40分。新聞40分。外出先はスーパー＞病院＞散歩。男性の15%が「2週間に1回以下」しか人と話さない。')

doc.add_heading('社会的接点マップ', level=3)
add_table(
    ['頻度', '接点', '備考'],
    [
        ['毎日〜週数回', 'テレビ（5時間40分）、新聞（40分）、スーパー（週3〜4回）', 'ほぼ全員が接触'],
        ['週1〜月数回', '病院・診療所（月2回以上が6割）、銀行・郵便局、散歩', '大半が接触'],
        ['月1回以下', '民生委員の訪問、地域包括の相談、老人クラブ、趣味サークル', '参加率は低下中'],
        ['ほぼ接点なし', '終活セミナー、弁護士事務所、葬儀社', '自ら行く動機がない'],
    ],
    col_widths=[1.3, 3.0, 2.2]
)

doc.add_heading('重要な発見：セミナーの限界', level=3)
add_para('「最も死後事務委任を必要としている層ほど、終活セミナーに来ない」', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))

add_table(
    ['セミナーに来る人', 'セミナーに来ない人'],
    [
        ['60代後半〜70代前半の女性', '75歳以上の男性'],
        ['まだ元気で外出できる', '外出が億劫・困難'],
        ['社交的・情報収集に積極的', '社会的に孤立'],
        ['「終活」を前向きに捉えている', '「自分の死」を考えたくない'],
    ],
    col_widths=[3.3, 3.3]
)

doc.add_heading('病院待合室のパンフレット設置：撤回', level=3)
add_para('当初、病院の待合室（月2回以上通院する高齢者が60%）をリーチポイントとして検討したが、以下の理由で撤回。')
add_bullet('病院側の心理：「死」を連想させるものを院内に置くことへの抵抗')
add_bullet('患者側の心理：体調が悪い状態で「死後の手続き」は不快')
add_bullet('医師の倫理観：特定の弁護士事務所の宣伝に加担することへの抵抗')

doc.add_page_break()

# =====================================================
# Phase 2
# =====================================================
doc.add_heading('Phase 2：戦略設計', level=1)

# --- 1. 預託金管理 ---
doc.add_heading('1. 預託金の管理方法', level=2)

doc.add_heading('依頼者名義の信託口座を弁護士が管理する方式の問題点', level=3)
add_table(
    ['問題', '内容'],
    [
        ['口座凍結リスク', '依頼者死亡の瞬間に口座凍結→死後事務の費用が引き出せない。本末転倒'],
        ['預金債権の帰属が曖昧', '倒産隔離効果が不確実'],
        ['弁護士会の想定外', '預り金口座は弁護士名義が前提（職務基本規程38条）'],
    ],
    col_widths=[2.0, 4.5]
)

doc.add_heading('推奨される3方式の比較', level=3)
add_table(
    ['方式', '安全性', 'コスト', '実務負担', 'おすすめ度'],
    [
        ['A. 弁護士預り金口座', '○', '無料', '低い', '◎ まず始めるならこれ'],
        ['B. 信託会社併用', '◎（二重の倒産隔離）', '口座開設費+保管料', 'やや高い', '○ 規模拡大後'],
        ['C. 遺産清算方式（預託金なし）', '◎', '公正証書遺言の作成費', '中程度', '◎ 初期費用を下げたいなら最適'],
    ],
    col_widths=[2.0, 1.0, 1.5, 1.0, 1.5]
)

doc.add_heading('遺産清算方式の詳細フロー', level=3)
add_para('【生前】', bold=True)
add_bullet('依頼者と弁護士が2つの公正証書を作成：①死後事務委任契約 ②公正証書遺言（弁護士を遺言執行者に指定）')
add_para('【死亡後】', bold=True)
add_bullet('Step 1：弁護士が死亡を確認')
add_bullet('Step 2：弁護士が「遺言執行者」として銀行に通知→遺言執行者は単独で口座解約・払戻しが可能（民法1012条2項）')
add_bullet('Step 3：遺産から死後事務費用（葬儀代・報酬等）を清算')
add_bullet('Step 4：残余があれば遺言の指定先へ')

add_para('遺産清算方式の課題：', bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
add_bullet('戸籍謄本（出生〜死亡の連続戸籍）の取得に1〜3週間かかる')
add_bullet('銀行の審査に1〜3週間かかる')
add_bullet('合計3〜6週間の間、葬儀費用（15〜50万円）の立替が不可避')
add_bullet('遺産が極端に少ない場合は費用回収できないリスク')

add_para('解決策：', bold=True)
add_bullet('「預託金方式」と「遺産清算方式」の選択制にする')
add_bullet('遺産清算方式の場合は少額預託金（葬儀費用分30〜50万円のみ）の併用を推奨')
add_bullet('葬儀保険の活用（保険金が約5営業日で下りる）')

# --- 2. 弁護士費用 ---
doc.add_heading('2. 弁護士費用の値付け', level=2)

doc.add_heading('業務別の原価構造', level=3)
add_table(
    ['業務', '弁護士の工数', '外注実費'],
    [
        ['契約書作成（公正証書）', '3〜5時間', '公証人手数料1.1〜4.3万'],
        ['死亡届提出', '1時間', '0'],
        ['葬儀手配・立会い', '3〜5時間', '葬儀代15〜50万（直葬なら15万〜）'],
        ['納骨手配', '2〜3時間', '永代供養3〜30万'],
        ['行政手続き一式', '3〜5時間', '0'],
        ['公共料金解約', '2〜3時間', '0'],
        ['賃貸明け渡し', '5〜10時間', '原状回復5〜30万'],
        ['遺品整理手配', '2〜3時間', '業者費用10〜40万'],
        ['デジタルアカウント削除', '2〜3時間', '0'],
        ['関係者への死亡通知', '2〜3時間', '郵送費数千円'],
    ],
    col_widths=[2.5, 1.5, 2.5]
)
add_note('弁護士の稼働合計：約25〜40時間')

doc.add_heading('競合との価格比較', level=3)
add_table(
    ['提供者', '契約時費用', '執行時報酬', '預託金', '合計'],
    [
        ['行政書士', '11〜33万', '22〜55万', '50〜100万', '80〜190万'],
        ['司法書士', '22〜33万', '33〜55万', '50〜100万', '100〜190万'],
        ['NPO（きずなの会等）', '44〜64万', '込み', '70〜130万', '120〜200万'],
        ['弁護士（大手）', '33〜55万', '55〜110万', '50〜150万', '140〜315万'],
    ],
    col_widths=[1.8, 1.2, 1.2, 1.2, 1.2]
)

doc.add_heading('提案：3段階プラン', level=3)

add_para('シンプルプラン（死後事務の最低限）', bold=True, size=Pt(11))
add_bullet('対象業務：死亡届・葬儀手配・納骨・行政手続き')
add_bullet('契約時：22万円（税込） / 執行時：33万円（税込） / 預託金：30〜50万円')
add_bullet('合計目安：85〜105万円。遺産清算方式なら初期費用22万円のみ')

add_para('スタンダードプラン（賃貸住まいの方向け）', bold=True, size=Pt(11))
add_bullet('対象業務：シンプル＋賃貸明渡し・遺品整理・公共料金解約・死亡通知')
add_bullet('契約時：33万円 / 執行時：55万円 / 預託金：50〜100万円')
add_bullet('合計目安：138〜188万円。遺産清算方式なら初期費用33万円のみ')

add_para('フルサポートプラン（全部おまかせ）', bold=True, size=Pt(11))
add_bullet('対象業務：スタンダード＋デジタル遺品・ペット・遺言執行')
add_bullet('契約時：44万円 / 執行時：88万円 / 預託金：80〜150万円')
add_bullet('合計目安：212〜282万円。遺産清算方式なら初期費用44万円のみ')

add_para('価格ポジショニング：', bold=True)
add_para('行政書士と同等〜やや上の価格帯で、弁護士品質を提供。「弁護士なのにこの価格」が最大の差別化ポイント。')

# --- 3. CPA分析 ---
doc.add_heading('3. 集客チャネルとCPA分析', level=2)

add_table(
    ['順位', 'チャネル', '推定CPA', '月間コスト', '高齢者リーチ', '安定性'],
    [
        ['1', 'セミナー（葬儀社共催）', '0.5〜2万円', '1〜3万', '◎ 直接接触', '△'],
        ['2', '地方紙小枠広告', '1〜6万円', '5〜20万', '◎ 購読率82.9%', '○'],
        ['3', 'ラジオ法律相談コーナー', 'ほぼ0円', '0〜3万', '○ 独居高齢者と相性◎', '◎'],
        ['4', '提携先紹介（賃貸管理等）', '1〜3万円', '2〜5万', '◎ 周囲経由', '△'],
        ['5', 'SEO', '1〜3万円', '0〜5万', '✕ 本人に届かない', '◎'],
        ['6', 'ポスティング（セミナー誘導）', '3〜8万円', '5〜10万', '○ 団地直配', '△'],
        ['7', 'リスティング広告', '2〜7万円', '6〜20万', '✕ 本人に届かない', '○'],
        ['8', 'ラジオCM（スポット）', '約64万円', '45万', '○', '○'],
    ],
    col_widths=[0.5, 2.0, 1.2, 1.2, 1.5, 0.6]
)

add_para('LTVとCACの関係：', bold=True)
add_bullet('スタンダードプラン基準LTV = 契約時33万＋執行時55万 = 88万円')
add_bullet('許容CAC = LTV ÷ 3 = 約29万円')
add_bullet('上記の全チャネル（ラジオスポットCM以外）は許容範囲内')

add_note('SEO・リスティングは高齢者本人より周囲（甥姪・ケアマネ等）が検索する前提。間接リーチとして有効。')

# --- 4. メッセージング ---
doc.add_heading('4. メッセージング戦略', level=2)

add_para('「死」を前面に出さない表現設計が極めて重要。', bold=True)

add_table(
    ['NG表現', 'OK表現'],
    [
        ['死後事務委任契約のご案内', 'おひとりさまの「もしも」に備える安心サポート'],
        ['あなたの死後の手続きを代行', '一人暮らしの「いざという時」、誰に頼みますか？'],
        ['遺品整理・葬儀の事前準備', 'あなたの暮らしを最後まで守る、弁護士の安心プラン'],
    ],
    col_widths=[3.3, 3.3]
)
add_para('サービスの本質は死後事務委任だが、入口のメッセージは「安心」「備え」「もしも」で包むこと。')

add_para('アクションの転換：', bold=True)
add_table(
    ['旧（NGパターン）', '新（推奨パターン）'],
    [
        ['チラシ→セミナーに来てください', 'チラシ→お電話ください（フリーダイヤル）'],
    ],
    col_widths=[3.3, 3.3]
)
add_para('独り身高齢者は外出が億劫。でも電話はかけられる。特に独り身は「誰かと話したい」欲求がある。電話相談をゴールにすることで、行動ハードルを大幅に下げる。')

# --- 5. チャネル設計 ---
doc.add_heading('5. 最終チャネル設計', level=2)

doc.add_heading('月3件プラン（月予算10〜20万円）', level=3)
add_table(
    ['チャネル', '月間コスト', '受任見込み', '備考'],
    [
        ['新聞折込チラシ（月2回）', '5〜15万円', '1〜2件', '「おひとりさまの安心相談」電話相談をゴールに'],
        ['ラジオ法律相談コーナー', '0〜3万円', '0.5件', '地方AM局に売り込み。パーソナリティとの対話形式'],
        ['賃貸管理会社提携（3〜5社）', '2〜3万円', '0.5〜1件', '営業経費のみ。入居条件として推奨'],
        ['信金連携セミナー（月1回）', '1〜2万円', '0.5〜1件', '「資産の引継ぎ」文脈。信金の会議室で開催'],
    ],
    col_widths=[2.0, 1.2, 1.2, 2.2]
)
add_para('合計：月8〜23万円、月2.5〜4.5件受任見込み', bold=True)
add_para('着手金30万円 × 3件 = 月90万円の売上に対し、広告費10〜20万円。ROAS 450〜900%')

doc.add_heading('月10件プラン（月予算27〜49万円）', level=3)
add_para('月3件プランに以下を追加：')
add_bullet('セミナー拡大（月4回に）：+3〜5万円、+2〜3件')
add_bullet('サ高住提携（10施設）：+3〜5万円、+1件')
add_bullet('居住支援法人連携：+1〜2万円、+1件')
add_bullet('SEOコンテンツ（月2〜4本）：+5〜10万円、+1〜2件（中期効果）')
add_bullet('ポスティング（セミナー誘導）：+5〜10万円、+0.5〜1件')

# --- 6. ロードマップ ---
doc.add_heading('6. 段階別ロードマップ', level=2)

doc.add_heading('Phase A：即日〜1ヶ月目（即効性重視）', level=3)
add_bullet('MEO（Googleマイビジネス）設定（無料・即日）')
add_bullet('弁護士ドットコム掲載（月2万円・3日で開始）')
add_bullet('リスティング広告開始（月5〜10万円・1週間で開始）')
add_bullet('葬儀社1社にセミナー共催を打診')
add_bullet('賃貸管理会社への営業開始（DM+電話）')
add_bullet('目標：月1〜2件受任')

doc.add_heading('Phase B：2〜6ヶ月目（提携網構築）', level=3)
add_bullet('葬儀社セミナー月2回を定常化')
add_bullet('賃貸管理会社3〜5社との提携確立')
add_bullet('居住支援法人との連携打診')
add_bullet('SEO記事投稿開始（月2本）')
add_bullet('自治体への終活セミナー講師登録')
add_bullet('地方ラジオ局に法律相談コーナー出演を売り込み')
add_bullet('目標：月3〜5件受任')

doc.add_heading('Phase C：7〜12ヶ月目（スケール）', level=3)
add_bullet('提携先を15〜20社に拡大')
add_bullet('SEOが上位表示され始め自然流入増')
add_bullet('セミナー月4回体制')
add_bullet('LINE公式アカウントでリピート・紹介の仕組み化')
add_bullet('新聞折込チラシを月2回で定常実施')
add_bullet('目標：月7〜10件受任')

# --- 7. 課題 ---
doc.add_heading('7. 課題・ネック・反対意見の整理', level=2)
add_para('本検討で浮上した全ての課題・反論を以下に整理する。事業化にあたり、各項目への対策が必要。', bold=True)

add_table(
    ['分類', '課題・ネック', '深刻度', '対策の方向性'],
    [
        ['法的制約', '個人情報保護法（リスト共有不可）', '高', '提携先の自社発送 or セミナー共催に切替'],
        ['法的制約', '弁護士法72条（紹介料支払不可）', '高', '課題解決ベースの関係構築に徹する'],
        ['法的制約', '弁護士広告規制（日弁連規程第6条）', '中', '特定事件の勧誘にならないメッセージ設計'],
        ['ターゲット', '独り身高齢者はセミナーに来ない可能性', '高', '電話相談をゴールに。「来てもらう」から「向こうから目に入る」設計へ'],
        ['ターゲット', '病院待合室は「死」の連想で設置困難', '高', '病院は撤回。新聞・ラジオ・信金等に集中'],
        ['ターゲット', 'SEO・リスティングは高齢者本人に届かない', '中', '周囲（甥姪・ケアマネ）経由の間接リーチとして活用'],
        ['競合', 'きずなの会が先行モデル（ブルーオーシャンではない）', '中', '死後事務委任特化、料金透明性、弁護士の直接担当で差別化'],
        ['収益', '遺産清算方式の立替リスク', '中', '少額預託金（葬儀費用分のみ）との選択制'],
        ['収益', '立ち上げ期は赤字〜トントン', '中', '3年目以降の執行報酬ストック収入で安定化'],
        ['運営', '弁護士1人の対応上限（月10件が限界）', '中', '事務局スタッフ1名の採用が必須'],
        ['市場', '新聞購読率の年7%減少トレンド', '低', '5〜10年は70代以上で有効。その間にSEO・LINEを育成'],
        ['メッセージ', '「死後事務委任」の表現が重い', '高', '「おひとりさまの安心サポート」等で包む'],
    ],
    col_widths=[1.0, 2.5, 0.7, 2.5]
)

doc.add_paragraph()
add_para('以上', alignment=WD_ALIGN_PARAGRAPH.RIGHT)

# ===== 保存 =====
output_path = r'C:\Users\user\Desktop\claudeマスター\死後事務委任_検討レポート.docx'
doc.save(output_path)
print(f'レポート生成完了: {output_path}')
